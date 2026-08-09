"""
DOCX parser using python-docx for extracting text from Word documents.
"""
from docx import Document


def parse_docx_file(file_path: str) -> str:
    """Extract text from all paragraphs and tables in a .docx file.

    Args:
        file_path: Absolute path to the .docx file.

    Returns:
        Concatenated text from all paragraphs and table cells.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If no text could be extracted.
    """
    doc = Document(file_path)
    parts = []

    # Extract paragraph text
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table text (clinical notes sometimes use tables)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    if not parts:
        raise ValueError(f"No text could be extracted from DOCX: {file_path}")

    return "\n".join(parts)
