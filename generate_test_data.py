"""
OmniFHIR-AI: Comprehensive Synthetic Test Data Generator

Generates 12 synthetic clinical document files across 4 categories and
6 modality types to exercise every feature of the pipeline:

Category A: Core Modality Coverage (Happy Path)
  1. patient_01_discharge.txt      — TXT, HbA1c 7.2%, COMPLIANT
  2. patient_02_chart.pdf          — PDF, HbA1c 9.4%, NON-COMPLIANT
  3. patient_03_labslip.png        — PNG, HbA1c 8.1%, NON-COMPLIANT (Vision OCR)
  4. patient_04_consult.docx       — DOCX, HbA1c 6.5%, COMPLIANT

Category B: OCR Stress Tests
  5. patient_05_fax_noisy.png      — Noisy PNG, HbA1c 7.8%, Tesseract fallback
  6. patient_06_handwritten.jpg    — Script font JPG, HbA1c 10.2%, Discrepancy flagging
  7. patient_07_multipage_scan.tiff— Multi-page TIFF, HbA1c 5.9%, COMPLIANT

Category C: Edge Cases
  8. patient_08_borderline.pdf     — PDF, HbA1c 8.0% (exact threshold) + old 7.5%
  9. patient_09_multiple_tests.docx— DOCX, Multi-test panel, HbA1c 6.8%
  10. patient_10_narrative_buried.txt— TXT, 500+ words, HbA1c 11.3% buried

Category D: Error Handling
  11. patient_11_corrupt.pdf       — Corrupt PDF
  12. patient_12_empty.txt         — Empty file
"""
import os
import random

from docx import Document
from PIL import Image, ImageDraw, ImageFont, ImageFilter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

DATA_DIR = "./sample_data"
os.makedirs(DATA_DIR, exist_ok=True)


# ═══════════════════════════════════════════════════════════════════════════════
# CATEGORY A: Core Modality Coverage (Happy Path)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_01_txt():
    """Patient 01: Plain text discharge note. HbA1c 7.2% — COMPLIANT."""
    content = """PATIENT VISIT NOTE
===================================================================

Patient Name: John Doe
Date of Birth: 1968-05-14
Patient ID: PT-10492
Date of Service: 2026-03-12
Provider: Dr. Amanda Chen, MD — Internal Medicine
Facility: Riverside Community Health Center

CHIEF COMPLAINT:
Routine diabetic follow-up appointment.

HISTORY OF PRESENT ILLNESS:
Mr. Doe is a 57-year-old male with a 12-year history of Type 2 Diabetes
Mellitus. He presents today for his scheduled quarterly diabetic evaluation.
Patient reports good medication adherence with Metformin 1000mg BID and
Glipizide 5mg daily. He denies any episodes of hypoglycemia, polyuria,
or polydipsia since his last visit. He has been monitoring his blood
glucose at home 2-3 times per week with readings ranging from 110-145 mg/dL
fasting.

PHYSICAL EXAMINATION:
- Vitals: BP 128/78, HR 72, Temp 98.4F, Weight 198 lbs
- General: Well-appearing, no acute distress
- Extremities: No peripheral edema, pedal pulses 2+ bilaterally
- Neurological: Monofilament testing normal bilateral feet

LABORATORY DATA:
Point-of-care Hemoglobin A1c test drawn today shows HbA1c of 7.2%.
Fasting glucose: 128 mg/dL.

ASSESSMENT:
Type 2 Diabetes Mellitus — well controlled on current regimen.

PLAN:
1. Continue Metformin 1000mg BID.
2. Continue Glipizide 5mg daily.
3. Annual dilated eye exam — due next month, referral sent.
4. Recheck HbA1c in 3 months.
5. Continue home glucose monitoring.

Electronically signed by Dr. Amanda Chen, MD
Date: 2026-03-12"""
    with open(os.path.join(DATA_DIR, "patient_01_discharge.txt"), "w", encoding="utf-8") as f:
        f.write(content.strip())
    print("  [OK] patient_01_discharge.txt -- TXT, HbA1c 7.2%, COMPLIANT")


def generate_02_pdf():
    """Patient 02: PDF discharge summary. HbA1c 9.4% — NON-COMPLIANT."""
    pdf_path = os.path.join(DATA_DIR, "patient_02_chart.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 750, "METRO HEALTH HOSPITAL")
    c.setFont("Helvetica", 10)
    c.drawString(72, 735, "Department of Endocrinology -- Discharge Summary")
    c.line(72, 730, 540, 730)
    c.setFont("Helvetica-Bold", 11)
    c.drawString(72, 710, "PATIENT INFORMATION")
    c.setFont("Helvetica", 10)
    info = [
        "Patient ID: PT-88201                    MRN: MH-20260218-001",
        "Patient Name: Maria Gonzalez            DOB: 1975-11-22",
        "Admission Date: 2026-02-15              Discharge Date: 2026-02-18",
        "Attending Physician: Dr. Sarah Jenkins, MD -- Endocrinology",
    ]
    y = 695
    for line in info:
        c.drawString(72, y, line)
        y -= 14
    y -= 10
    c.setFont("Helvetica-Bold", 11)
    c.drawString(72, y, "PRINCIPAL DIAGNOSIS")
    y -= 14
    c.setFont("Helvetica", 10)
    c.drawString(72, y, "Type 2 Diabetes Mellitus with Poor Glycemic Control (ICD-10: E11.65)")
    y -= 24
    c.setFont("Helvetica-Bold", 11)
    c.drawString(72, y, "HOSPITAL COURSE")
    y -= 14
    c.setFont("Helvetica", 10)
    for line in [
        "Patient presented to the ED on 02/15 with altered mental status, severe",
        "dehydration, and serum glucose of 485 mg/dL. She was admitted for IV fluid",
        "resuscitation and insulin drip protocol. Transitioned to SubQ insulin on day 2.",
    ]:
        c.drawString(72, y, line)
        y -= 14
    y -= 10
    c.setFont("Helvetica-Bold", 11)
    c.drawString(72, y, "DISCHARGE LABORATORY RESULTS")
    y -= 14
    c.setFont("Helvetica", 10)
    for line in [
        "Serum Glucose (fasting): 210 mg/dL         [Reference: 70-100 mg/dL]  HIGH",
        "Hemoglobin A1c: 9.4%                        [Reference: < 5.7%]        HIGH",
        "Serum Creatinine: 1.1 mg/dL                 [Reference: 0.6-1.2 mg/dL] Normal",
    ]:
        c.drawString(72, y, line)
        y -= 14
    y -= 10
    c.setFont("Helvetica-Bold", 11)
    c.drawString(72, y, "DISCHARGE RECOMMENDATIONS")
    y -= 14
    c.setFont("Helvetica", 10)
    for line in [
        "1. Titrate basal insulin (Lantus) to 30 units nightly.",
        "2. Refer to certified diabetes educator within 2 weeks.",
        "3. Follow-up with Dr. Jenkins in 1 week.",
        "4. Recheck HbA1c in 3 months -- target < 8.0%.",
    ]:
        c.drawString(72, y, line)
        y -= 14
    c.save()
    print("  [OK] patient_02_chart.pdf -- PDF, HbA1c 9.4%, NON-COMPLIANT")


def generate_03_png():
    """Patient 03: Clean lab result slip image. HbA1c 8.1% — NON-COMPLIANT."""
    img = Image.new("RGB", (700, 350), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("C:/Windows/Fonts/consola.ttf", 14)
    except (OSError, IOError):
        font = ImageFont.load_default()
    lab_text = [
        "QUEST DIAGNOSTICS -- LABORATORY REPORT",
        "=" * 50,
        "PATIENT: James Wilson    PATIENT ID: PT-55012",
        "DOB: 1960-08-30         COLLECTION DATE: 2026-04-02",
        "ORDERING PHYSICIAN: Dr. Robert Park, MD",
        "=" * 50,
        "TEST NAME              RESULT    REF RANGE   FLAG",
        "-" * 50,
        "Hemoglobin A1c         8.1 %     4.0-5.6 %   HIGH",
        "Glucose, Fasting       142 mg/dL 70-100       HIGH",
        "BUN                    18 mg/dL  7-20         Normal",
        "Creatinine             0.9 mg/dL 0.6-1.2      Normal",
        "-" * 50,
        "INTERPRETATION: HbA1c indicates UNCONTROLLED diabetes",
    ]
    y = 15
    for line in lab_text:
        d.text((20, y), line, fill=(0, 0, 0), font=font)
        y += 22
    img.save(os.path.join(DATA_DIR, "patient_03_labslip.png"))
    print("  [OK] patient_03_labslip.png -- PNG, HbA1c 8.1%, NON-COMPLIANT")


def generate_04_docx():
    """Patient 04: Endo consult note. HbA1c 6.5% — COMPLIANT."""
    doc = Document()
    doc.add_heading("ENDOCRINOLOGY CONSULTATION NOTE", level=1)
    doc.add_paragraph(
        "Patient ID: PT-30219 | Date of Assessment: 2026-01-25 | "
        "Referring Physician: Dr. Karen White, MD"
    )
    doc.add_heading("Reason for Referral", level=2)
    doc.add_paragraph(
        "Patient referred for comprehensive diabetes management evaluation. "
        "55-year-old female with Type 2 Diabetes diagnosed in 2018."
    )
    doc.add_heading("Assessment & Findings", level=2)
    p = doc.add_paragraph()
    p.add_run("Laboratory Results (collected 2026-01-24):\n").bold = True
    p.add_run(
        "Recent routine lab panel confirms Hemoglobin A1c level is 6.5% as of "
        "2026-01-24. This represents excellent glycemic control."
    )
    doc.add_heading("Plan", level=2)
    doc.add_paragraph(
        "1. Continue current Metformin regimen.\n"
        "2. Follow-up in 6 months with repeat HbA1c."
    )
    doc.save(os.path.join(DATA_DIR, "patient_04_consult.docx"))
    print("  [OK] patient_04_consult.docx -- DOCX, HbA1c 6.5%, COMPLIANT")


# ═══════════════════════════════════════════════════════════════════════════════
# CATEGORY B: OCR Stress Tests
# ═══════════════════════════════════════════════════════════════════════════════

def generate_05_noisy_png():
    """Patient 05: Noisy faxed lab result. HbA1c 7.8% — triggers Tesseract fallback."""
    img = Image.new("RGB", (650, 300), color=(245, 242, 238))
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("C:/Windows/Fonts/consola.ttf", 13)
    except (OSError, IOError):
        font = ImageFont.load_default()
    d.text((10, 5), "FAX TRANSMISSION: 2026-04-15 14:32 EST  PG 1/1", fill=(80, 80, 80), font=font)
    d.line((10, 22, 640, 22), fill=(100, 100, 100), width=1)
    lab_lines = [
        "LABCORP - CLINICAL LABORATORY REPORT",
        "",
        "PATIENT: Robert Kim         ID: PT-67210",
        "DOB: 1972-03-19             COLLECTED: 2026-04-14",
        "PHYSICIAN: Dr. Lisa Chang",
        "",
        "TEST                 RESULT     REFERENCE",
        "--------------------------------------------",
        "Hemoglobin A1c       7.8 %      4.0 - 5.6 %",
        "Glucose, Random      156 mg/dL  70 - 140",
        "",
        "** END OF REPORT **",
    ]
    y = 35
    for line in lab_lines:
        d.text((25, y), line, fill=(30, 30, 30), font=font)
        y += 20
    # Add speckle noise
    pixels = img.load()
    for _ in range(800):
        x = random.randint(0, 649)
        yy = random.randint(0, 299)
        gray = random.randint(100, 180)
        pixels[x, yy] = (gray, gray, gray)
    img = img.filter(ImageFilter.GaussianBlur(radius=0.8))
    img = img.rotate(1.5, expand=True, fillcolor=(245, 242, 238))
    img.save(os.path.join(DATA_DIR, "patient_05_fax_noisy.png"))
    print("  [OK] patient_05_fax_noisy.png -- Noisy PNG, HbA1c 7.8%, Tesseract fallback")


def generate_06_handwritten_jpg():
    """Patient 06: Simulated handwritten note. HbA1c 10.2% — discrepancy flagging."""
    img = Image.new("RGB", (700, 350), color=(255, 253, 245))
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("C:/Windows/Fonts/comic.ttf", 16)
    except (OSError, IOError):
        font = ImageFont.load_default()
    lines = [
        "Clinical Note - Dr. Michael Torres",
        "Date: 2026-05-08",
        "",
        "Pt: Sandra Hughes  ID: PT-91403",
        "DOB: 1965-12-01",
        "",
        "Dx: T2DM - poorly controlled",
        "Labs drawn 5/7/2026:",
        "  HbA1c = 10.2%  (very high!)",
        "  FBG = 245 mg/dL",
        "",
        "Plan: Start insulin glargine 20u QHS",
        "  Increase metformin to 1000mg BID",
        "  F/u 2 weeks",
    ]
    y = 15
    for line in lines:
        x_off = random.randint(-2, 2)
        y_off = random.randint(-1, 1)
        d.text((25 + x_off, y + y_off), line, fill=(20, 20, 80), font=font)
        y += 22
    img.save(os.path.join(DATA_DIR, "patient_06_handwritten.jpg"), quality=85)
    print("  [OK] patient_06_handwritten.jpg -- Script JPG, HbA1c 10.2%, Discrepancy flag")


def generate_07_multipage_tiff():
    """Patient 07: Multi-page TIFF scan. HbA1c 5.9% on page 2 — COMPLIANT."""
    try:
        font = ImageFont.truetype("C:/Windows/Fonts/consola.ttf", 13)
        font_bold = ImageFont.truetype("C:/Windows/Fonts/consolab.ttf", 15)
    except (OSError, IOError):
        font = ImageFont.load_default()
        font_bold = font
    # Page 1: Demographics
    page1 = Image.new("RGB", (700, 400), color=(255, 255, 255))
    d1 = ImageDraw.Draw(page1)
    d1.text((25, 20), "CITYWIDE MEDICAL CENTER", fill=(0, 0, 0), font=font_bold)
    d1.text((25, 45), "PATIENT DEMOGRAPHICS", fill=(0, 0, 0), font=font)
    d1.line((25, 65, 675, 65), fill=(0, 0, 0), width=2)
    for i, line in enumerate([
        "Patient ID: PT-44829",
        "Name: Catherine Bell",
        "DOB: 1980-06-15  Age: 45  Sex: Female",
        "PCP: Dr. William Foster, MD",
        "",
        "*** SEE PAGE 2 FOR LABORATORY RESULTS ***",
    ]):
        d1.text((25, 80 + i * 18), line, fill=(0, 0, 0), font=font)
    # Page 2: Lab results
    page2 = Image.new("RGB", (700, 400), color=(255, 255, 255))
    d2 = ImageDraw.Draw(page2)
    d2.text((25, 20), "CITYWIDE MEDICAL CENTER -- LAB RESULTS", fill=(0, 0, 0), font=font_bold)
    d2.text((25, 42), "Patient: Catherine Bell (PT-44829)  Date: 2026-06-10", fill=(0, 0, 0), font=font)
    d2.line((25, 60, 675, 60), fill=(0, 0, 0), width=2)
    for i, line in enumerate([
        "TEST                    RESULT      REFERENCE",
        "=" * 50,
        "Glucose, Fasting        102 mg/dL   70-100",
        "Hemoglobin A1c          5.9 %       4.0-5.6 %",
        "Creatinine              0.8 mg/dL   0.6-1.2",
        "Total Cholesterol       195 mg/dL   <200",
    ]):
        d2.text((25, 75 + i * 17), line, fill=(0, 0, 0), font=font)
    tiff_path = os.path.join(DATA_DIR, "patient_07_multipage_scan.tiff")
    page1.save(tiff_path, save_all=True, append_images=[page2])
    print("  [OK] patient_07_multipage_scan.tiff -- Multi-page TIFF, HbA1c 5.9%, COMPLIANT")


# ═══════════════════════════════════════════════════════════════════════════════
# CATEGORY C: Edge Cases & Confidence Challenges
# ═══════════════════════════════════════════════════════════════════════════════

def generate_08_borderline_pdf():
    """Patient 08: Borderline HbA1c 8.0% + old value 7.5%."""
    pdf_path = os.path.join(DATA_DIR, "patient_08_borderline.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, 750, "PRIMARY CARE CLINIC -- FOLLOW-UP NOTE")
    c.setFont("Helvetica", 10)
    c.drawString(72, 730, "Patient ID: PT-40021 | Date: 2026-07-15")
    c.line(72, 725, 540, 725)
    lines = [
        "Mr. David Park is a 62-year-old male returning for diabetes follow-up.",
        "He was seen 6 months ago in January 2026 when his HbA1c was 7.5%.",
        "At that time, we increased his Metformin and added Sitagliptin.",
        "",
        "Unfortunately, medication adherence has been inconsistent.",
        "",
        "CURRENT LABORATORY RESULTS (drawn 2026-07-14):",
        "  Hemoglobin A1c: 8.0%  (Previous: 7.5% on 2026-01-10)",
        "  Fasting Glucose: 165 mg/dL",
        "",
        "ASSESSMENT: The current HbA1c of 8.0% is at the HEDIS threshold.",
        "The January value of 7.5% showed better control but the current",
        "reading is what matters for compliance reporting.",
    ]
    y = 710
    for line in lines:
        c.drawString(72, y, line)
        y -= 14
    c.save()
    print("  [OK] patient_08_borderline.pdf -- PDF, HbA1c 8.0% (borderline), NON-COMPLIANT")


def generate_09_multiple_tests_docx():
    """Patient 09: Multi-test lab panel with HbA1c buried among other results."""
    doc = Document()
    doc.add_heading("COMPREHENSIVE METABOLIC & DIABETES PANEL", level=1)
    doc.add_paragraph(
        "Patient ID: PT-72019 | Name: Emily Chen | DOB: 1978-09-03\n"
        "Collection Date: 2026-05-20"
    )
    doc.add_heading("Complete Blood Count (CBC)", level=2)
    table1 = doc.add_table(rows=4, cols=4)
    table1.style = "Table Grid"
    for i, h in enumerate(["Test", "Result", "Reference", "Flag"]):
        table1.rows[0].cells[i].text = h
    for r, data in enumerate([
        ("WBC", "7.2 K/uL", "4.5-11.0", "Normal"),
        ("Hemoglobin", "14.1 g/dL", "12.0-16.0", "Normal"),
        ("Platelets", "245 K/uL", "150-400", "Normal"),
    ], 1):
        for i, v in enumerate(data):
            table1.rows[r].cells[i].text = v
    doc.add_heading("Lipid Panel", level=2)
    table2 = doc.add_table(rows=4, cols=4)
    table2.style = "Table Grid"
    for i, h in enumerate(["Test", "Result", "Reference", "Flag"]):
        table2.rows[0].cells[i].text = h
    for r, data in enumerate([
        ("Total Cholesterol", "210 mg/dL", "<200", "Borderline"),
        ("LDL", "128 mg/dL", "<130", "Near Optimal"),
        ("Triglycerides", "150 mg/dL", "<150", "Borderline"),
    ], 1):
        for i, v in enumerate(data):
            table2.rows[r].cells[i].text = v
    doc.add_heading("Diabetes Markers", level=2)
    table3 = doc.add_table(rows=3, cols=4)
    table3.style = "Table Grid"
    for i, h in enumerate(["Test", "Result", "Reference", "Flag"]):
        table3.rows[0].cells[i].text = h
    for r, data in enumerate([
        ("Glucose, Fasting", "118 mg/dL", "70-100", "HIGH"),
        ("Hemoglobin A1c", "6.8%", "4.0-5.6%", "Elevated"),
    ], 1):
        for i, v in enumerate(data):
            table3.rows[r].cells[i].text = v
    doc.add_paragraph(
        "\nInterpretation: HbA1c of 6.8% indicates near-target glycemic control."
    )
    doc.save(os.path.join(DATA_DIR, "patient_09_multiple_tests.docx"))
    print("  [OK] patient_09_multiple_tests.docx -- DOCX, Multi-test, HbA1c 6.8%, COMPLIANT")


def generate_10_narrative_buried():
    """Patient 10: Long narrative with HbA1c 11.3% buried deep."""
    content = """COMPREHENSIVE ANNUAL WELLNESS VISIT -- PHYSICIAN NARRATIVE

Patient ID: PT-23891
Patient Name: Thomas Martinez
Date of Birth: 1958-02-28
Date of Visit: 2026-06-05
Provider: Dr. Elizabeth Warren, MD -- Family Medicine

SOCIAL HISTORY UPDATE:
Mr. Martinez is a retired postal worker who lives with his wife of 38 years
in a single-story home. He reports he has been mostly sedentary since
retirement two years ago, spending most of his time gardening and watching
television. He denies tobacco use (quit 15 years ago after a 20-pack-year
history). He reports occasional alcohol consumption, approximately 2-3 beers
per week on weekends. His diet consists primarily of home-cooked meals
prepared by his wife, though he admits to frequent snacking between meals.

REVIEW OF SYSTEMS:
Constitutional: Patient reports feeling more tired than usual over the
past 3 months. No significant weight change reported. No fevers or chills.
HEENT: No vision changes reported, last eye exam over 18 months ago.
Cardiovascular: No chest pain, palpitations, or orthopnea. Reports
occasional mild bilateral lower extremity edema. Respiratory: No shortness
of breath at rest. Reports mild dyspnea on exertion. GI: Reports increased
thirst and frequent urination over the past several weeks which is concerning
in the context of his diabetes history. Musculoskeletal: Bilateral knee pain
rated 4/10. Reports morning stiffness lasting approximately 20 minutes.
Neurological: Reports intermittent tingling in bilateral feet, present for
approximately 6 months and gradually worsening.

CURRENT MEDICATIONS:
1. Metformin 1000mg BID (for Type 2 Diabetes, diagnosed 2014)
2. Glipizide 10mg BID (added 2022 due to worsening control)
3. Lisinopril 20mg daily (Hypertension)
4. Amlodipine 5mg daily (Hypertension, added 2023)
5. Atorvastatin 40mg daily (Hyperlipidemia)
6. Aspirin 81mg daily (Cardiovascular prophylaxis)
7. Omeprazole 20mg daily (GERD)
8. Gabapentin 300mg TID (Diabetic peripheral neuropathy, started 2025)

LABORATORY RESULTS:
Today's laboratory work was drawn fasting this morning. His complete
metabolic panel shows sodium of 141, potassium 4.3, chloride 101, CO2
of 25, BUN 22 slightly elevated, creatinine 1.3 up from 1.1 six months
ago worth monitoring, and eGFR of 58 which is now in Stage 3a CKD
territory. Liver function tests within normal limits. His lipid panel shows
total cholesterol of 198, LDL of 112, HDL of 42, and triglycerides of 220
which are elevated. Of significant concern is his Hemoglobin A1c which has
risen substantially to 11.3%, up from 8.8% six months ago. This represents
a dramatic worsening of his glycemic control despite being on maximum doses
of two oral agents. His fasting glucose today was 268 mg/dL confirming the
poor control indicated by the HbA1c result.

ASSESSMENT AND PLAN:
1. Type 2 Diabetes Mellitus -- POORLY CONTROLLED (HbA1c 11.3%)
   Initiating basal insulin therapy with Lantus 15 units at bedtime.
   Continue Metformin, discontinue Glipizide.
2. Diabetic Peripheral Neuropathy -- PROGRESSIVE
3. Chronic Kidney Disease Stage 3a -- NEW DIAGNOSIS
4. Hypertension -- CONTROLLED
5. Hyperlipidemia -- LDL near goal, triglycerides elevated.

Follow-up in 2 weeks for insulin titration.

Electronically signed: Dr. Elizabeth Warren, MD
Date: 2026-06-05"""
    with open(os.path.join(DATA_DIR, "patient_10_narrative_buried.txt"), "w", encoding="utf-8") as f:
        f.write(content.strip())
    print("  [OK] patient_10_narrative_buried.txt -- TXT, 500+ words, HbA1c 11.3% buried")


# ═══════════════════════════════════════════════════════════════════════════════
# CATEGORY D: Error Handling & Graceful Degradation
# ═══════════════════════════════════════════════════════════════════════════════

def generate_11_corrupt_pdf():
    """Patient 11: Intentionally corrupted PDF."""
    pdf_path = os.path.join(DATA_DIR, "patient_11_corrupt.pdf")
    with open(pdf_path, "wb") as f:
        f.write(b"%PDF-1.4\n")
        f.write(b"1 0 obj\n")
        f.write(b"<< /Type /Catalog >>\n")
        f.write(b"endobj\n")
        f.write(b"%%EOF\n")
        f.write(bytes(random.getrandbits(8) for _ in range(100)))
    print("  [OK] patient_11_corrupt.pdf -- Corrupt PDF, SHOULD FAIL at ingestion")


def generate_12_empty_txt():
    """Patient 12: Completely empty file."""
    with open(os.path.join(DATA_DIR, "patient_12_empty.txt"), "w", encoding="utf-8") as f:
        pass
    print("  [OK] patient_12_empty.txt -- Empty file (0 bytes), SHOULD FAIL")


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("  OmniFHIR-AI: Synthetic Test Data Generator")
    print("=" * 60)
    print()
    print("Category A: Core Modality Coverage")
    print("-" * 55)
    generate_01_txt()
    generate_02_pdf()
    generate_03_png()
    generate_04_docx()
    print()
    print("Category B: OCR Stress Tests")
    print("-" * 55)
    generate_05_noisy_png()
    generate_06_handwritten_jpg()
    generate_07_multipage_tiff()
    print()
    print("Category C: Edge Cases")
    print("-" * 55)
    generate_08_borderline_pdf()
    generate_09_multiple_tests_docx()
    generate_10_narrative_buried()
    print()
    print("Category D: Error Handling")
    print("-" * 55)
    generate_11_corrupt_pdf()
    generate_12_empty_txt()
    print()
    print("=" * 60)
    print(f"  Generated 12 test files in '{DATA_DIR}/'")
    print()
    summary = [
        ("patient_01_discharge.txt",       ".txt",  "7.2%",  "COMPLIANT",     "Text parser"),
        ("patient_02_chart.pdf",           ".pdf",  "9.4%",  "NON-COMPLIANT", "PDF parser"),
        ("patient_03_labslip.png",         ".png",  "8.1%",  "NON-COMPLIANT", "Vision OCR"),
        ("patient_04_consult.docx",        ".docx", "6.5%",  "COMPLIANT",     "DOCX parser"),
        ("patient_05_fax_noisy.png",       ".png",  "7.8%",  "COMPLIANT",     "Tesseract fallback"),
        ("patient_06_handwritten.jpg",     ".jpg",  "10.2%", "NON-COMPLIANT", "Discrepancy flag"),
        ("patient_07_multipage_scan.tiff", ".tiff", "5.9%",  "COMPLIANT",     "Multi-page TIFF"),
        ("patient_08_borderline.pdf",      ".pdf",  "8.0%",  "NON-COMPLIANT", "Temporal disambig."),
        ("patient_09_multiple_tests.docx", ".docx", "6.8%",  "COMPLIANT",     "Multi-test panel"),
        ("patient_10_narrative_buried.txt", ".txt", "11.3%", "NON-COMPLIANT", "Needle-in-haystack"),
        ("patient_11_corrupt.pdf",         ".pdf",  "N/A",   "FAIL (Stage 1)","Error handling"),
        ("patient_12_empty.txt",           ".txt",  "N/A",   "FAIL (Stage 3)","Empty input"),
    ]
    header = f"{'Filename':<38} {'Type':<7} {'HbA1c':<7} {'Expected':<16} {'Tests'}"
    print(header)
    print("-" * len(header))
    for row in summary:
        print(f"{row[0]:<38} {row[1]:<7} {row[2]:<7} {row[3]:<16} {row[4]}")


if __name__ == "__main__":
    main()